"""
Nome do Script: docx_export.py
Autor: Renato Borges
Data: 22 de Março de 2026
Versão: 2.1.0
Propósito: Exportação do conteúdo final revisado para o formato Microsoft Word (DOCX).
           Refatorado com foco rigoroso em Mobile-First (18pt) e layout BNCC.
"""

import logging
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional

# Configuração de Logging Profissional para rastreabilidade
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("DocxExporter")

class DocxExporter:
    """
    Especialista em converter o conteúdo processado pelo Ebook Generator
    em documentos Word formatados profissionalmente para leitura em smartphones.
    """

    def __init__(self):
        self.document = Document()

    def _configurar_estilos_padrao(self):
        """Configura fontes, tamanhos (18pt) e margens (Mobile-First) para o documento."""
        # Configuração de Layout Mobile (A5/Smartphone)
        section = self.document.sections[0]
        section.page_width = Cm(14.8)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(1.5)

        # Estilo Normal: Corpo do Texto - 18pt
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(18)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Estilo Heading 1: Títulos de Capítulo - 28pt
        h1 = self.document.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(28)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(31, 56, 100) # Azul BNCC

        # Estilo Heading 2: Subseções/Perguntas - 22pt
        h2 = self.document.styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(22)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(64, 64, 64) # Cinza Grafite

    def gerar_ebook_docx(self, 
                         titulo: str, 
                         autor: str, 
                         conteudo: str, 
                         output_path: str = "ebook_final.docx",
                         subtitulo: Optional[str] = None) -> Optional[str]:
        """
        Cria o arquivo DOCX estruturado com capa (texto) e corpo do manuscrito.
        """
        try:
            logger.info(f"Iniciando exportação DOCX Mobile-First: {output_path}")
            self._configurar_estilos_padrao()

            # 1. Capa Interna (Texto Centralizado)
            # Espaçamento inicial para simular centralização vertical
            for _ in range(5): self.document.add_paragraph()
            
            title_part = self.document.add_heading(titulo.upper(), level=1)
            title_part.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if subtitulo:
                sub = self.document.add_paragraph(subtitulo)
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Aplicando itálico via run para o subtítulo
                if sub.runs: sub.runs[0].italic = True

            author_para = self.document.add_paragraph(f"\n\nAutor: {autor}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self.document.add_page_break()

            # 2. Processamento de Conteúdo BNCC
            paragrafos = conteudo.split('\n')
            for p in paragrafos:
                texto_limpo = p.strip()
                if texto_limpo:
                    # Lógica de detecção de Títulos (H1 se for curto e uppercase)
                    if texto_limpo.isupper() and len(texto_limpo) < 60:
                        self.document.add_heading(texto_limpo, level=1)
                    # Lógica para perguntas ou subseções (H2)
                    elif texto_limpo.endswith('?') or (texto_limpo.startswith('Dúvida') and len(texto_limpo) < 100):
                        self.document.add_heading(texto_limpo, level=2)
                    else:
                        self.document.add_paragraph(texto_limpo)

            self.document.save(output_path)
            logger.info(f"✅ Sucesso: {output_path} gerado com fonte 18pt.")
            return output_path

        except Exception as e:
            logger.error(f"❌ Falha crítica na exportação: {e}")
            return None

# Autor: Renato Borges
