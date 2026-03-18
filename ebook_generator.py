"""
Nome do Script: ebook_designer_pro.py
Autor: Renato Borges
Data: 18 de Março de 2026
Versão: 2.0.0
Propósito: Gerar eBook .docx com layout profissional, tipografia avançada 
           e estruturação baseada no Documento Curricular de Goiás (DCGO).
"""

import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuração de Logging para excelência operacional
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class EbookArchitect:
    """
    Motor de diagramação focado em estética e hierarquia visual.
    """

    def __init__(self, main_title: str, subtitle: str):
        self.doc = Document()
        self.main_title = main_title
        self.subtitle = subtitle
        logging.info(f"Iniciando Engine de Design para: {self.main_title}")

    def _apply_font_settings(self, style, name: str, size: int, color_rgb: tuple, bold: bool = False):
        """
        Aplica configurações refinadas de fonte e cor aos estilos.
        """
        font = style.font
        font.name = name
        font.size = Pt(size)
        font.bold = bold
        font.color.rgb = RGBColor(*color_rgb)
        
        # Garante a aplicação da fonte no motor de renderização do Word
        r_fonts = style.element.rPr.get_or_add_rFonts()
        r_fonts.set(qn('w:eastAsia'), name)

    def setup_professional_styles(self):
        """
        Configura os estilos conforme a teoria das cores e tipografia solicitada.
        """
        # H1 - Título de Capítulo (28pt, Azul Escuro/Slate)
        h1 = self.doc.styles['Heading 1']
        self._apply_font_settings(h1, 'Segoe UI', 28, (31, 56, 100), True)
        
        # H2 - Subtítulos (22pt, Cinza Profundo)
        h2 = self.doc.styles['Heading 2']
        self._apply_font_settings(h2, 'Segoe UI', 22, (64, 64, 64), True)
        
        # Texto Normal - Corpo do Ebook (18pt para acessibilidade e leitura fluida)
        normal = self.doc.styles['Normal']
        self._apply_font_settings(normal, 'Calibri', 18, (0, 0, 0), False)
        
        # Espaçamento entre parágrafos para evitar "paredes de texto"
        normal.paragraph_format.space_after = Pt(14)
        normal.paragraph_format.line_spacing = 1.15
        
        logging.info("Paleta de cores e tipografia injetadas no documento.")

    def insert_page_numbers(self):
        """
        Insere numeração de páginas automática no rodapé para navegação profissional.
        """
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        def add_field(run, field_type):
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = field_type
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        run = p.add_run("Página ")
        add_field(run, "PAGE")

    def create_layout(self, content_blocks: list, output_name: str):
        """
        Monta o documento respeitando as quebras de página e hierarquia.
        """
        try:
            self.setup_professional_styles()
            
            # --- Capa ---
            capa = self.doc.add_paragraph()
            capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = capa.add_run(f"{self.main_title}\n")
            run_title.font.size = Pt(36)
            run_title.bold = True
            
            run_sub = capa.add_run(self.subtitle)
            run_sub.font.size = Pt(20)
            self.doc.add_page_break()

            # --- Loop de Conteúdo ---
            for block in content_blocks:
                if block['style'] == 'H1':
                    self.doc.add_heading(block['text'], level=1)
                elif block['style'] == 'H2':
                    self.doc.add_heading(block['text'], level=2)
                else:
                    p = self.doc.add_paragraph(block['text'])
                    p.alignment = WD_ALIGN_PARAGRAPH.BOTH

                if block.get('new_page'):
                    self.doc.add_page_break()

            self.insert_page_numbers()
            self.doc.save(output_name)
            logging.info(f"Sucesso! Documento '{output_name}' gerado com layout premium.")
            
        except Exception as e:
            logging.error(f"Falha na geração do layout: {e}")

if __name__ == "__main__":
    # Dados extraídos do documento modelo [cite: 7, 12, 13]
    conteudo_refinado = [
        {"style": "H1", "text": "INTRODUÇÃO", "new_page": False},
        {"style": "Para", "text": "A inserção da Computação na Educação Básica é uma das maiores inovações curriculares do nosso tempo...", "new_page": True},
        {"style": "H1", "text": "CAPÍTULO 1: IMPLEMENTAÇÃO", "new_page": False},
        {"style": "H2", "text": "A escola precisa criar uma nova disciplina?", "new_page": False},
        {"style": "Para", "text": "Não. A orientação é trabalhar a Computação integrada às disciplinas que já existem de forma transversal.", "new_page": False},
    ]

    # Execução do robô
    designer = EbookArchitect("BNCC DA COMPUTAÇÃO", "Perguntas Frequentes & Guia Prático")
    designer.create_layout(conteudo_refinado, "Ebook_BNCC_Computacao_Final.docx")
