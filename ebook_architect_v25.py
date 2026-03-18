"""
Nome do Script: ebook_architect_v25.py
Autor: Renato Borges
Data: 18 de Março de 2026
Versão: 2.5.0 (Estável - Anti-502)
Propósito: Diagramação profissional BNCC com otimização de performance.
"""

import logging
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

# Configuração de Logging para auditoria
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class EbookRobustEngine:
    def __init__(self, title: str):
        self.doc = Document()
        self.title = title
        self._setup_layout()

    def _setup_layout(self):
        """Configura o layout Mobile-First e Estilos de forma otimizada."""
        section = self.doc.sections[0]
        section.page_width = Cm(14.8) # Largura Mobile
        
        # Estilo Normal - 18pt
        style_n = self.doc.styles['Normal']
        style_n.font.name = 'Calibri'
        style_n.font.size = Pt(18)
        style_n.paragraph_format.line_spacing = 1.5
        
        # Heading 1 - 28pt (Azul BNCC)
        h1 = self.doc.styles['Heading 1']
        h1.font.size = Pt(28)
        h1.font.color.rgb = RGBColor(31, 56, 100)
        
        # Heading 2 - 22pt (Cinza)
        h2 = self.doc.styles['Heading 2']
        h2.font.size = Pt(22)
        h2.font.color.rgb = RGBColor(64, 64, 64)

    def _add_page_number_lite(self):
        """Versão simplificada para evitar sobrecarga de XML."""
        footer = self.doc.sections[0].footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        
        def fld(name):
            e = OxmlElement(f'w:{name}')
            if name == 'fldChar': e.set(ns.qn('w:fldCharType'), 'begin')
            return e

        # Inserção direta de campo de página
        run._r.append(fld('fldChar'))
        instr = OxmlElement('w:instrText'); instr.text = "PAGE"
        run._r.append(instr)
        end = OxmlElement('w:fldChar'); end.set(ns.qn('w:fldCharType'), 'end')
        run._r.append(end)

    def build(self, content: list, filename: str):
        """Build com tratamento de exceção e logging por etapa."""
        try:
            # Título Inicial
            logging.info("Criando Capa...")
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"\n\n{self.title.upper()}")
            r.bold = True
            r.font.size = Pt(32)
            self.doc.add_page_break()

            # Processamento do Conteúdo (Otimizado)
            for i, block in enumerate(content):
                tag, text = block.get("type"), block.get("text", "")
                logging.info(f"Processando bloco {i+1}...")
                
                if tag == "h1":
                    self.doc.add_heading(text, level=1)
                elif tag == "h2":
                    self.doc.add_heading(text, level=2)
                else:
                    self.doc.add_paragraph(text)

            self._add_page_number_lite()
            self.doc.save(filename)
            logging.info(f"✅ Sucesso: {filename} gerado.")
            
        except Exception as e:
            logging.error(f"❌ Erro no build: {e}")

if __name__ == "__main__":
    # Exemplo de dados para teste rápido
    data_test = [
        {"type": "h1", "text": "BNCC DA COMPUTAÇÃO"},
        {"type": "p", "text": "Este manual organiza as principais dúvidas sobre o DC-GO."},
        {"type": "h2", "text": "Dúvida 01: Transversalidade"},
        {"type": "p", "text": "A computação deve ser integrada e não uma disciplina isolada."}
    ]
    
    app = EbookRobustEngine("Guia Prático Professor")
    app.build(data_test, "Ebook_BNCC_Final.docx")
